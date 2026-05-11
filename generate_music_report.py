#!/usr/bin/env python3
"""
🎵 Music Veille Agent
Agent IA pour suivre les nouveautés musicales françaises hebdomadairement.
Powered by Claude + GitHub Actions
"""

import os
import sys
from datetime import datetime, timedelta
from pathlib import Path

from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CONFIG
# ============================================================================

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not ANTHROPIC_API_KEY:
    print("❌ ERREUR: ANTHROPIC_API_KEY non trouvée")
    sys.exit(1)

REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

# ============================================================================
# AGENT CLASS
# ============================================================================

class MusicVeilleAgent:
    """Agent IA pour la veille musicale française."""
    
    def __init__(self):
        self.client = Anthropic()
        self.model = "claude-opus-4-6"
        self.conversation = []
        self.week_start = (datetime.now() - timedelta(days=datetime.now().weekday())).strftime("%d/%m/%Y")
        self.report_data = {}
    
    def add_message(self, role: str, content: str):
        """Ajouter un message à la conversation."""
        self.conversation.append({
            "role": role,
            "content": content
        })
    
    def get_response(self, max_tokens: int = 2000) -> str:
        """Appeler Claude et retourner la réponse."""
        print(f"📡 Appel Claude API ({len(self.conversation)} messages)...")
        
        response = self.client.messages.create(
            model=self.model,
            max_tokens=max_tokens,
            system=self._get_system_prompt(),
            messages=self.conversation
        )
        
        content = response.content[0].text
        self.conversation.append({
            "role": "assistant",
            "content": content
        })
        
        return content
    
    def _get_system_prompt(self) -> str:
        """Prompt système de l'agent."""
        return """Tu es un expert passionné en musique française 2026.
Tu dois faire une veille hebdomadaire complète et de qualité sur les sorties musicales francophones.

Domaines à couvrir:
- Variété française
- Pop Variété
- Chanson française (tradition + moderne)
- Nouvelle scène (artistes émergents, indie)
- Hip-hop/rap français
- Électro/synthwave française

Style:
- Sois enthousiaste mais professionnel
- Cite toujours tes sources (Spotify, YouTube, Pitchfork, etc.)
- Mentionne les artistes établis ET les nouveaux talents
- Explique le contexte/story derrière les sorties
- Ajoute des recommandations/playlists

Format:
- Clair et bien structuré
- Sections distinctes
- Bulletpoints quand utile
- Pas trop long (vise 1500-2000 mots max)"""
    
    # ========================================================================
    # BOUCLE D'AGENT (Multi-turn conversation)
    # ========================================================================
    
    def step_1_research(self):
        """ÉTAPE 1: Recherche complète des sorties musiques."""
        print("\n🔍 ÉTAPE 1: Recherche des sorties musicales...")
        
        prompt = f"""Semaine du {self.week_start}

Fais une recherche complète sur les NOUVELLES SORTIES MUSICALES FRANCOPHONES cette semaine.

Inclus:
1. Les 10 meilleures nouvelles sorties (singles, albums, clips)
2. Artistes qui font buzz
3. Collaborations intéressantes
4. Récompenses/news industrie
5. Tendances observées

Sois détaillé et spécifique (noms d'artistes, titres, dates, labels)."""
        
        self.add_message("user", prompt)
        research = self.get_response(max_tokens=2500)
        self.report_data["research"] = research
        
        print("✅ Recherche complète")
        return research
    
    def step_2_analysis(self):
        """ÉTAPE 2: Analyse et synthèse intelligente."""
        print("\n📊 ÉTAPE 2: Analyse et synthèse...")
        
        prompt = """Basé sur ta recherche précédente, réponds:

1. Quel est le THÈME DOMINANT cette semaine?
2. Y a-t-il un artiste qui se démarque particulièrement?
3. Quelles tendances musicales observes-tu?
4. Quel est ton TOP 3 PERSONNEL (justifié)?

Sois concis mais insightful (max 500 mots)."""
        
        self.add_message("user", prompt)
        analysis = self.get_response(max_tokens=1500)
        self.report_data["analysis"] = analysis
        
        print("✅ Analyse complète")
        return analysis
    
    def step_3_recommendations(self):
        """ÉTAPE 3: Recommandations et playlists."""
        print("\n🎧 ÉTAPE 3: Recommandations...")
        
        prompt = """Suggère:

1. Une PLAYLIST de 8-10 chansons pour découvrir la semaine (mix sorties + classiques)
2. Un artiste ÉMERGENT à suivre absolument
3. UN PODCAST/INTERVIEW à écouter (musicien ou producteur français)

Format:
- Playlist: [Titre] - Artiste (année/label)
- Émergent: [Nom] - Description courte
- Podcast: [Titre] - Host/Artiste"""
        
        self.add_message("user", prompt)
        recommendations = self.get_response(max_tokens=1000)
        self.report_data["recommendations"] = recommendations
        
        print("✅ Recommandations générées")
        return recommendations
    
    def step_4_export_word(self):
        """ÉTAPE 4: Exporter en document Word professionnel."""
        print("\n📄 ÉTAPE 4: Export Word...")
        
        doc = Document()
        
        # Header
        title = doc.add_heading("🎵 VEILLE MUSIQUE FRANÇAISE", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"Semaine du {self.week_start}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(14)
        subtitle_format.font.italic = True
        
        doc.add_paragraph()  # Espace
        
        # Recherche
        doc.add_heading("📻 SORTIES & ACTUALITÉS", level=1)
        doc.add_paragraph(self.report_data.get("research", ""))
        
        # Analyse
        doc.add_page_break()
        doc.add_heading("🎯 ANALYSE DE LA SEMAINE", level=1)
        doc.add_paragraph(self.report_data.get("analysis", ""))
        
        # Recommandations
        doc.add_page_break()
        doc.add_heading("🎧 À ÉCOUTER CETTE SEMAINE", level=1)
        doc.add_paragraph(self.report_data.get("recommendations", ""))
        
        # Footer
        doc.add_page_break()
        footer_text = doc.add_paragraph()
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_text.add_run(
            f"Rapport généré par 🤖 Music Veille Agent\n"
            f"Claude + GitHub Actions\n"
            f"{datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Sauvegarder
        filename = REPORTS_DIR / f"veille_{datetime.now().strftime('%Y-%m-%d')}.docx"
        doc.save(filename)
        print(f"✅ Document sauvegardé: {filename}")
        
        return str(filename)
    
    def run(self):
        """Lance l'agent complet."""
        print("=" * 70)
        print("🎵 MUSIC VEILLE AGENT - DÉMARRAGE")
        print("=" * 70)
        
        try:
            # Exécuter la boucle d'agent
            self.step_1_research()
            self.step_2_analysis()
            self.step_3_recommendations()
            report_path = self.step_4_export_word()
            
            print("\n" + "=" * 70)
            print("✅ SUCCÈS: Rapport généré avec succès!")
            print(f"📁 Fichier: {report_path}")
            print("=" * 70)
            
            return True
            
        except Exception as e:
            print(f"\n❌ ERREUR: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    agent = MusicVeilleAgent()
    success = agent.run()
    sys.exit(0 if success else 1)
